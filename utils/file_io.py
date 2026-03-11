"""
File I/O utilities for reading and writing scripts
Supports: .txt, .pdf, .doc, .docx
"""

import os
import json
from pathlib import Path
from config import (
    RAW_SCRIPTS_DIR,
    CLEANED_SCRIPTS_DIR,
    SEGMENTED_SCRIPTS_DIR,
    OUTPUT_DIR
)

# Import document processing libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

if not PDFPLUMBER_AVAILABLE and not PDF_AVAILABLE:
    print("Warning: Neither pdfplumber nor PyPDF2 installed. PDF support disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled.")

def ensure_directories():
    """
    Create necessary directories if they don't exist
    """
    directories = [
        RAW_SCRIPTS_DIR,
        CLEANED_SCRIPTS_DIR,
        SEGMENTED_SCRIPTS_DIR,
        OUTPUT_DIR
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)

def read_script(filepath):
    """
    Read script file with automatic format detection
    Supports: .txt, .pdf, .doc, .docx
    
    Args:
        filepath (str): Path to script file
        
    Returns:
        str: Script content
        
    Raises:
        FileNotFoundError: If file doesn't exist
        IOError: If file cannot be read
        ValueError: If file format is not supported
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Script file not found: {filepath}")
    
    # Get file extension
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    
    # Route to appropriate reader
    if ext == '.txt':
        return _read_txt(filepath)
    elif ext == '.pdf':
        return _read_pdf(filepath)
    elif ext in ['.doc', '.docx']:
        return _read_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .txt, .pdf, .doc, .docx")

def _read_txt(filepath):
    """
    Read plain text file with automatic encoding detection
    
    Args:
        filepath (str): Path to text file
        
    Returns:
        str: Text content
    """
    # Try different encodings
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            return content
        except UnicodeDecodeError:
            continue
    
    raise IOError(f"Could not decode text file with any standard encoding: {filepath}")

def _normalize_pdf_text(text):
    """
    Fix common PDF extraction issues for screenplays.
    
    PyPDF2 and other extractors often merge scene markers into
    the middle of lines. This function inserts line breaks before
    known scene markers so the rule-based segmenter can find them.
    """
    import re
    
    # Fix split words (e.g., "IN T." → "INT.", "EX T." → "EXT.")
    text = re.sub(r'IN\s+T\.', 'INT.', text)
    text = re.sub(r'EX\s+T\.', 'EXT.', text)
    
    # Insert newline before scene markers found mid-line
    text = re.sub(r'(?<!\n)(?<!^)(\s*(?:INT\.|EXT\.))', r'\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'(?<!\n)(?<!^)(\s*CUT\s+TO\s*:)', r'\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'(?<!\n)(?<!^)(\s*FADE\s+(?:IN|OUT)[.:])', r'\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'(?<!\n)(?<!^)(\s*FADE\s+TO\s+BLACK)', r'\n\1', text, flags=re.IGNORECASE)
    
    # Normalize excessive blank lines (more than 3 → 2)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    return text


def _read_pdf(filepath):
    """
    Read PDF file and extract text.
    
    Uses pdfplumber (better layout preservation) if available,
    falls back to PyPDF2.
    
    Args:
        filepath (str): Path to PDF file
        
    Returns:
        str: Extracted text content
        
    Raises:
        RuntimeError: If no PDF library is available
        IOError: If PDF cannot be read
    """
    if not PDFPLUMBER_AVAILABLE and not PDF_AVAILABLE:
        raise RuntimeError(
            "PDF support not available. Install with: pip install pdfplumber"
        )
    
    raw_text = None
    
    # Strategy 1: pdfplumber (better layout preservation)
    if PDFPLUMBER_AVAILABLE:
        try:
            text_content = []
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(text)
                    except Exception as e:
                        print(f"Warning: pdfplumber could not extract page {page_num + 1}: {e}")
                        continue
            
            if text_content:
                raw_text = "\n\n".join(text_content)
        except Exception as e:
            print(f"Warning: pdfplumber failed, trying PyPDF2: {e}")
            raw_text = None
    
    # Strategy 2: PyPDF2 fallback
    if raw_text is None and PDF_AVAILABLE:
        try:
            reader = PdfReader(filepath)
            text_content = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(text)
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            if text_content:
                raw_text = "\n\n".join(text_content)
        except Exception as e:
            raise IOError(f"Error reading PDF file: {e}")
    
    if not raw_text:
        raise IOError("No text could be extracted from PDF")
    
    # Normalize the extracted text to fix common PDF issues
    return _normalize_pdf_text(raw_text)

def _read_docx(filepath):
    """
    Read DOCX (or DOC if converted) file and extract text
    
    Args:
        filepath (str): Path to DOCX file
        
    Returns:
        str: Extracted text content
        
    Raises:
        RuntimeError: If DOCX library not available
        IOError: If DOCX cannot be read
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "DOCX support not available. Install with: pip install python-docx"
        )
    
    _, ext = os.path.splitext(filepath)
    if ext.lower() == '.doc':
        raise ValueError(
            "Legacy .doc format not directly supported. "
            "Please convert to .docx or use LibreOffice/Word to save as .docx"
        )
    
    try:
        doc = Document(filepath)
        text_content = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_content.append(text)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        if not text_content:
            raise IOError("No text could be extracted from DOCX")
        
        # Join all paragraphs with newlines
        return "\n".join(text_content)
    
    except Exception as e:
        raise IOError(f"Error reading DOCX file: {e}")

def detect_file_format(filepath):
    """
    Detect file format and check if it's supported
    
    Args:
        filepath (str): Path to file
        
    Returns:
        dict: Format information
    """
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    
    format_info = {
        "extension": ext,
        "supported": ext in ['.txt', '.pdf', '.docx'],
        "requires_library": None
    }
    
    if ext == '.pdf' and not PDF_AVAILABLE:
        format_info["supported"] = False
        format_info["requires_library"] = "PyPDF2"
    elif ext in ['.doc', '.docx']:
        if ext == '.doc':
            format_info["supported"] = False
            format_info["note"] = "Convert to .docx format"
        elif not DOCX_AVAILABLE:
            format_info["supported"] = False
            format_info["requires_library"] = "python-docx"
    
    return format_info

def save_output(data, filename, output_dir=OUTPUT_DIR):
    """
    Save processed output to JSON file
    
    Args:
        data (dict): Data to save
        filename (str): Output filename
        output_dir (str): Output directory
        
    Returns:
        str: Path to saved file
    """
    ensure_directories()
    
    filepath = os.path.join(output_dir, filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return filepath

def save_intermediate(content, filename, stage="cleaned"):
    """
    Save intermediate processing results
    
    Args:
        content (str or dict): Content to save
        filename (str): Output filename
        stage (str): Processing stage ('cleaned', 'segmented')
        
    Returns:
        str: Path to saved file
    """
    ensure_directories()
    
    # Determine directory based on stage
    if stage == "cleaned":
        directory = CLEANED_SCRIPTS_DIR
    elif stage == "segmented":
        directory = SEGMENTED_SCRIPTS_DIR
    else:
        directory = OUTPUT_DIR
    
    filepath = os.path.join(directory, filename)
    
    # Save based on content type
    if isinstance(content, dict) or isinstance(content, list):
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(content, f, indent=2, ensure_ascii=False)
    else:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(str(content))
    
    return filepath

def list_scripts(directory=RAW_SCRIPTS_DIR, extension=None):
    """
    List all script files in a directory
    
    Args:
        directory (str): Directory to search
        extension (str): File extension filter (e.g., '.txt', '.pdf')
        
    Returns:
        list: List of script file paths
    """
    if not os.path.exists(directory):
        return []
    
    files = []
    supported_extensions = ['.txt', '.pdf', '.docx']
    
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        
        if os.path.isfile(filepath):
            file_ext = os.path.splitext(filename)[1].lower()
            
            if extension is None:
                # No filter, but only include supported formats
                if file_ext in supported_extensions:
                    files.append(filepath)
            elif filename.endswith(extension):
                files.append(filepath)
    
    return sorted(files)

def get_output_path(input_filepath, suffix="_processed"):
    """
    Generate output filepath based on input filepath
    
    Args:
        input_filepath (str): Input file path
        suffix (str): Suffix to add before extension
        
    Returns:
        str: Output file path
    """
    ensure_directories()
    
    # Get filename without path and extension
    basename = os.path.basename(input_filepath)
    name_without_ext = os.path.splitext(basename)[0]
    
    # Create output filename (always JSON)
    output_filename = f"{name_without_ext}{suffix}.json"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    return output_path

def get_file_size(filepath):
    """
    Get file size in human-readable format
    
    Args:
        filepath (str): File path
        
    Returns:
        str: File size (e.g., "1.5 MB")
    """
    if not os.path.exists(filepath):
        return "0 B"
    
    size_bytes = os.path.getsize(filepath)
    
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    
    return f"{size_bytes:.1f} TB"

def get_file_info(filepath):
    """
    Get comprehensive file information
    
    Args:
        filepath (str): File path
        
    Returns:
        dict: File information
    """
    if not os.path.exists(filepath):
        return {"exists": False}
    
    stat = os.stat(filepath)
    _, ext = os.path.splitext(filepath)
    
    return {
        "exists": True,
        "path": filepath,
        "name": os.path.basename(filepath),
        "extension": ext.lower(),
        "size": get_file_size(filepath),
        "size_bytes": stat.st_size,
        "format_info": detect_file_format(filepath)
    }